"""Export Word cu python-docx — pagini landscape, tabel pe toată lățimea paginii."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from ui.export.export_html import group_spans
from ui.export.export_utils import format_cell_value, format_total_value, validate_pages

# Fonturi — mai mari decât înainte, lizibile la print
FONT_HEADER = 8.5
FONT_DATA = 9.5
FONT_TOTAL = 9.5


def export_to_word(out_path: Path, pages: list[dict]) -> Path:
    validate_pages(pages)
    doc = Document()
    _configure_section(doc.sections[0])

    for pi, page in enumerate(pages):
        if pi > 0:
            doc.add_page_break()
        if page.get("type") == "cover":
            _render_cover(doc, page)
        else:
            _render_page(doc, page)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def _configure_section(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Mm(8)
    section.right_margin = Mm(8)
    section.top_margin = Mm(8)
    section.bottom_margin = Mm(12)
    _add_page_number_footer(section)


def _usable_width(doc) -> int:
    section = doc.sections[-1]
    return int(section.page_width - section.left_margin - section.right_margin)


def _set_table_full_width(table, ncols: int, doc) -> None:
    """Tabel 100% lățime pagină, coloane egale — ca în previzualizarea din app."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    for tag in ("w:tblW", "w:tblLayout"):
        for el in tblPr.findall(qn(tag)):
            tblPr.remove(el)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")
    tblPr.append(tblW)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    if ncols > 0:
        col_w = _usable_width(doc) // ncols
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_w

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO


def _render_cover(doc, page: dict) -> None:
    def add(text, size, space_before=0):
        if not text:
            return
        if space_before:
            doc.add_paragraph()
        p = doc.add_paragraph(str(text))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(size)

    for _ in range(6):
        doc.add_paragraph()
    add(page.get("institutie_1"), 15)
    add(page.get("institutie_2"), 15)
    add(page.get("titlu"), 26, space_before=1)
    add(page.get("biblioteca"), 22)
    add(page.get("localitate"), 17, space_before=1)
    add(page.get("an"), 19, space_before=1)


def _render_page(doc, page: dict) -> None:
    headers = page["headers"]
    groups = page["groups"]
    col_keys = page["col_keys"]
    rows = page["rows"]
    total_rows = page["total_rows"]
    meta = page["meta"]

    if meta.get("nume_biblioteca"):
        loc = f", {meta['localitate']}" if meta.get("localitate") else ""
        p = doc.add_paragraph(f"{meta['nume_biblioteca']}{loc}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)

    p = doc.add_paragraph("Registru de evidență a activității bibliotecii")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(13)

    partea = f"Partea {meta.get('parte_roman', '')}. {meta.get('title', '')}"
    if meta.get("luna_name"):
        partea += f" în luna {meta['luna_name']} anul {meta.get('an', '')}"
    elif meta.get("an"):
        partea += f" — anul {meta.get('an', '')}"
    p = doc.add_paragraph(partea)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(11)

    ncols = len(headers)
    has_groups = any(groups)
    header_rows = 2 if has_groups else 1
    total_count = len(total_rows)

    table = doc.add_table(rows=header_rows + len(rows) + total_count, cols=ncols)
    table.style = "Table Grid"
    table.autofit = False

    spans = group_spans(groups)
    if has_groups:
        grp_row = table.rows[0]
        lbl_row = table.rows[1]
        for start, end, g in spans:
            if g:
                merged = grp_row.cells[start]
                for c in range(start + 1, end + 1):
                    merged = merged.merge(grp_row.cells[c])
                merged.text = g
                _style_cell(merged, bold=True, size_pt=FONT_HEADER)
                for c in range(start, end + 1):
                    lbl_row.cells[c].text = headers[c]
                    _style_cell(lbl_row.cells[c], bold=True, size_pt=FONT_HEADER)
            else:
                cell = grp_row.cells[start].merge(lbl_row.cells[start])
                cell.text = headers[start]
                _style_cell(cell, bold=True, size_pt=FONT_HEADER)
    else:
        for c, h in enumerate(headers):
            table.rows[0].cells[c].text = h
            _style_cell(table.rows[0].cells[c], bold=True, size_pt=FONT_HEADER)

    base = header_rows
    for ri, row in enumerate(rows):
        for ci, key in enumerate(col_keys):
            val = format_cell_value(row.get(key, ""))
            cell = table.rows[base + ri].cells[ci]
            cell.text = val
            _style_cell(cell, bold=(ci == 0), size_pt=FONT_DATA)

    base = header_rows + len(rows)
    for ti, (label, sums) in enumerate(total_rows):
        rr = table.rows[base + ti]
        rr.cells[0].text = label
        _style_cell(rr.cells[0], bold=True, size_pt=FONT_TOTAL)
        for ci, key in enumerate(col_keys):
            if ci == 0:
                continue
            val = format_total_value(sums.get(key, ""))
            rr.cells[ci].text = val
            _style_cell(rr.cells[ci], bold=True, size_pt=FONT_TOTAL)

    _set_table_full_width(table, ncols, doc)


def _add_page_number_footer(section) -> None:
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run("Pagina ")
    _add_field(run, "PAGE")


def _add_field(run, field_code: str) -> None:
    fldChar1 = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    instrText = run._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instrText.text = f" {field_code} "
    fldChar2 = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _style_cell(cell, *, bold: bool = False, size_pt: float = FONT_DATA) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for par in cell.paragraphs:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = par.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        if not par.runs:
            par.add_run(par.text or "")
        for run in par.runs:
            run.bold = bold
            run.font.size = Pt(size_pt)
